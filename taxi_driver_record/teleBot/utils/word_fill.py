from docx import Document
import os

from .dbcommands import get_user_info, get_user_info_for_word_script, get_phone_nums_of_relative


async def word_filling(tele_id):

    info = await get_user_info_for_word_script(tele_id)
    info_relative_phones = await get_phone_nums_of_relative(tele_id)
    info.update(info_relative_phones)
    variables = info
    template_file_path = 'example.docx'
    output_file_path = f'WordFiles/{variables["${pk}"]}.docx'


    template_document = Document(template_file_path)

    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        for table in template_document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, variable_key, variable_value)

    template_document.save(output_file_path)


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)